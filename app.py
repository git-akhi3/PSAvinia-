import streamlit as st
import pandas as pd
from io import BytesIO
from utils import generate_summary, export_summary
from scraping import scrape_google_scholar
from docx import Document

st.set_page_config(page_title="Publications Summary Generator", layout="wide")

st.title("📚 Publications Summary Generator")
st.write("Easily generate summaries of faculty publications for profiles and reports.")

# Sidebar for file upload and options
st.sidebar.header("Upload and Options")

uploaded_file = st.sidebar.file_uploader("Upload Excel File", type=["xlsx"])

if uploaded_file:
    try:
        df = pd.read_excel(uploaded_file)
        st.success("File uploaded successfully!")
    except Exception as e:
        st.error(f"Error reading the file: {e}")
        st.stop()
else:
    st.warning("Please upload an Excel file to proceed.")
    st.stop()

st.header("Publication Data Preview")
st.dataframe(df)

# Filter Options
st.sidebar.subheader("Filter Publications")

authors = df['Author Name'].unique()
selected_authors = st.sidebar.multiselect("Select Authors", options=authors, default=authors)

years = df['Year'].unique()
selected_years = st.sidebar.multiselect("Select Years", options=years, default=years)

pub_types = df['Publication Type'].unique()
selected_types = st.sidebar.multiselect("Select Publication Types", options=pub_types, default=pub_types)

# Apply Filters
filtered_df = df[
    (df['Author Name'].isin(selected_authors)) &
    (df['Year'].isin(selected_years)) &
    (df['Publication Type'].isin(selected_types))
]

st.header("Filtered Publication Data")
st.dataframe(filtered_df)

st.sidebar.subheader("Generate Summary")

summary_format = st.sidebar.selectbox("Select Output Format", options=["Word Document", "Excel File"])

if st.sidebar.button("Generate Summary"):
    if filtered_df.empty:
        st.error("No data available to generate summary. Adjust your filters.")
    else:
        if summary_format == "Word Document":
            # Generate Word Document
            doc = Document()
            doc.add_heading('Publication Summary', 0)

            for author in selected_authors:
                author_df = filtered_df[filtered_df['Author Name'] == author]
                if not author_df.empty:
                    doc.add_heading(author, level=1)
                    for index, row in author_df.iterrows():
                        p = doc.add_paragraph()
                        p.add_run(f"Title: ").bold = True
                        p.add_run(f"{row['Title']}\n")
                        p.add_run(f"Year: ").bold = True
                        p.add_run(f"{row['Year']}\n")
                        p.add_run(f"Publication Type: ").bold = True
                        p.add_run(f"{row['Publication Type']}\n")
                        p.add_run(f"Journal/Conference Name: ").bold = True
                        p.add_run(f"{row['Journal/Conference Name']}\n")

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="📥 Download Summary as Word Document",
                data=buffer,
                file_name="publication_summary.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.success("Word Document generated successfully!")

        elif summary_format == "Excel File":
            output = BytesIO()
            filtered_df.to_excel(output, index=False)
            output.seek(0)

            st.download_button(
                label="📥 Download Summary as Excel File",
                data=output,
                file_name="publication_summary.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
            st.success("Excel File generated successfully!")
